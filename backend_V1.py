import os
import time
import torch
import re
import pandas as pd
from typing import Tuple, List, Dict
import fitz
from docx import Document as DocxDocument
from llama_index.core import Document, Settings, VectorStoreIndex
from llama_index.llms.huggingface import HuggingFaceLLM
from llama_index.embeddings.langchain import LangchainEmbedding
from langchain_huggingface import HuggingFaceEmbeddings
from huggingface_hub import login
from transformers import BitsAndBytesConfig

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from threading import Lock

FILES_DIR = "files"
FILE_LOADERS = {
    '.txt': lambda p: open(p, 'r', encoding='utf-8').read(),
    '.pdf': lambda p: "".join(page.get_text() for page in fitz.open(p)),
    '.docx': lambda p: "\n".join(para.text for para in DocxDocument(p).paragraphs),
    '.csv': lambda p: pd.read_csv(p).to_string(index=False),
    '.xls': lambda p: pd.read_excel(p).to_string(index=False),
    '.xlsx': lambda p: pd.read_excel(p).to_string(index=False),
}

documents = []
query_engine = None
file_snapshot = {}
lock = Lock()

def hf_login(token: str):
    login(token=token)

def get_embed_model():
    return LangchainEmbedding(
        HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    )

def get_llm(system_prompt: str):
    quant_config = BitsAndBytesConfig(load_in_8bit=True)
    return HuggingFaceLLM(
        context_window=4096,
        max_new_tokens=256,
        system_prompt=system_prompt,
        tokenizer_name="meta-llama/Llama-3.2-1B-Instruct",
        model_name="meta-llama/Llama-3.2-1B-Instruct",
        device_map="auto",
        model_kwargs={
            "torch_dtype": torch.float16,
            "quantization_config": quant_config,
            "pad_token_id": 128001
        }
    )

def load_documents(directory: str) -> List[Document]:
    docs = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        ext = os.path.splitext(filename)[1].lower()
        if ext in FILE_LOADERS:
            try:
                content = FILE_LOADERS[ext](file_path)
                docs.append(Document(text=content, metadata={"file_name": filename}))
                print(f"✅ Loaded: {filename}")
            except Exception as e:
                print(f"❌ Failed to load {filename}: {e}")
        else:
            print(f"⚠️ Skipped unsupported file: {filename}")
    return docs

def build_index_and_engine(documents):
    index = VectorStoreIndex.from_documents(documents)
    return index, index.as_query_engine()

def get_files_snapshot(directory: str) -> Dict[str, float]:
    snapshot = {}
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            snapshot[filename] = os.path.getmtime(file_path)
    return snapshot

def parse_query(user_query: str) -> Tuple[str | None, str]:
    match = re.search(r'from\s+(file|document)\s+[\'"]?([^\'"]+\.\w+)[\'"]?', user_query, re.IGNORECASE)
    if match:
        file_name = match.group(2).strip()
        stripped_query = re.sub(r'from\s+(file|document)\s+[\'"]?([^\'"]+\.\w+)[\'"]?', '', user_query, flags=re.IGNORECASE).strip()
        return file_name, stripped_query
    return None, user_query

class FileChangeHandler(FileSystemEventHandler):
    def on_any_event(self, event):
        global documents, query_engine, file_snapshot
        if event.is_directory:
            return
        with lock:
            print("\n🔁 File change detected. Reloading documents...")
            documents = load_documents(FILES_DIR)
            if documents:
                _, query_engine = build_index_and_engine(documents)
                file_snapshot = get_files_snapshot(FILES_DIR)
                print("✅ Documents reloaded.\n")
            else:
                print("⚠️ No supported documents found after reload.")

def main():
    global documents, query_engine, file_snapshot

    if not os.path.exists(FILES_DIR):
        os.makedirs(FILES_DIR)
        print(f"📂 Created '{FILES_DIR}' folder. Please add some documents and restart.")
        return

    hf_login(token="")  # "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"

    system_prompt = (
        "You are a Q&A assistant. Your goal is to answer questions based on the given documents."
    )
    embed_model = get_embed_model()
    llm = get_llm(system_prompt)

    Settings.llm = llm
    Settings.embed_model = embed_model
    Settings.chunk_size = 1024
    Settings.context_window = 4096

    print("\n📚 Checking documents in 'files/' directory...")
    documents = load_documents(FILES_DIR)
    if not documents:
        print("⚠️ No supported documents found in 'files/' folder. Add .txt, .pdf, .docx, .csv, .xls, or .xlsx files.")
        return

    _, query_engine = build_index_and_engine(documents)
    file_snapshot = get_files_snapshot(FILES_DIR)

    # Start file monitoring
    event_handler = FileChangeHandler()
    observer = Observer()
    observer.schedule(event_handler, path=FILES_DIR, recursive=False)
    observer.start()

    print("\n🤖 Q&A Assistant is ready! Type your question (or 'exit' to quit):")
    try:
        while True:
            user_query = input("\n\n").strip()
            if user_query.lower() in ['exit', 'quit']:
                print("👋 Goodbye!")
                break

            with lock:
                target_file, stripped_query = parse_query(user_query)

                if target_file:
                    matching_docs = [doc for doc in documents if doc.metadata.get("file_name", "").lower() == target_file.lower()]
                    if matching_docs:
                        _, filtered_engine = build_index_and_engine(matching_docs)
                        response = filtered_engine.query(stripped_query)
                    else:
                        response = f"⚠️ No document named '{target_file}' found in 'files/' directory."
                else:
                    response = query_engine.query(user_query)

            print(f"Assistant: {response}\n")

    finally:
        observer.stop()
        observer.join()

if __name__ == "__main__":
    main()
