import os
import time
import torch
import re
import pandas as pd
from typing import Tuple, List, Dict
import fitz
from docx import Document as DocxDocument
from llama_index.core import Document, Settings, VectorStoreIndex
from llama_index.llms.huggingface import HuggingFaceLLM
from llama_index.embeddings.langchain import LangchainEmbedding
from langchain_huggingface import HuggingFaceEmbeddings
from huggingface_hub import login
from transformers import BitsAndBytesConfig
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from threading import Lock
import hashlib

from llama_index.core import StorageContext, load_index_from_storage
PERSIST_DIR = "index_storage"
FILES_DIR = "files"

def dataframe_to_markdown(df: pd.DataFrame, filename: str) -> str:
    MAX_ROWS = 50
    preview_df = df.head(MAX_ROWS)

    metadata = [
        f"### Table Preview from `{filename}`",
        f"- **Shape**: {df.shape[0]} rows × {df.shape[1]} columns",
        "- **Columns**:",
        *(f"  - `{col}`: {dtype}" for col, dtype in zip(df.columns, df.dtypes)),
        "",
        preview_df.to_markdown(index=False)
    ]
    return "\n".join(metadata)

FILE_LOADERS = {
    '.txt': lambda p: open(p, 'r', encoding='utf-8').read(),
    '.pdf': lambda p: "".join(page.get_text() for page in fitz.open(p)),
    '.docx': lambda p: "\n".join(para.text for para in DocxDocument(p).paragraphs),
    '.csv': lambda p: dataframe_to_markdown(pd.read_csv(p), os.path.basename(p)),
    '.xls': lambda p: dataframe_to_markdown(pd.read_excel(p), os.path.basename(p)),
    '.xlsx': lambda p: dataframe_to_markdown(pd.read_excel(p), os.path.basename(p)),
}

documents = []
query_engine = None
file_snapshot = {}
lock = Lock()

def hf_login(token: str):
    login(token=token)

def get_embed_model():
    return LangchainEmbedding(
        HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    )

def get_llm(system_prompt: str):
    quant_config = BitsAndBytesConfig(load_in_8bit=True)
    return HuggingFaceLLM(
        context_window=4096,
        max_new_tokens=256,
        system_prompt=system_prompt,
        tokenizer_name="meta-llama/Llama-3.2-1B-Instruct",
        model_name="meta-llama/Llama-3.2-1B-Instruct",
        device_map="auto",
        model_kwargs={
            "torch_dtype": torch.float16,
            "quantization_config": quant_config,
            "pad_token_id": 128001
        }
    )

from llama_index.core.node_parser import SimpleNodeParser

def chunk_documents(documents: List[Document]) -> List[Document]:
    parser = SimpleNodeParser.from_defaults(chunk_size=1024, chunk_overlap=200)
    nodes = []
    for doc in documents:
        try:
            nodes.extend(parser.get_nodes_from_documents([doc]))
        except Exception as e:
            print(f"❌ Failed to chunk document {doc.metadata.get('file_name', '')}: {e}")
    return nodes


def load_documents(directory: str) -> List[Document]:
    docs = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        ext = os.path.splitext(filename)[1].lower()
        if ext in FILE_LOADERS:
            try:
                content = FILE_LOADERS[ext](file_path)
                docs.append(Document(text=content, metadata={"file_name": filename}))
                print(f"✅ Loaded: {filename}")
            except Exception as e:
                print(f"❌ Failed to load {filename}: {e}")
        else:
            print(f"⚠️ Skipped unsupported file: {filename}")
    
    return chunk_documents(docs)


from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.core.schema import TextNode

def build_index_and_engine(documents: List[Document], persist: bool = True):
    # Step 1: Chunk documents into nodes
    nodes = chunk_documents(documents)

    # Step 2: Create a fresh storage context (do not load from disk)
    storage_context = StorageContext.from_defaults()

    # Step 3: Use Settings (already set globally) to build an empty index
    index = VectorStoreIndex.from_documents([], storage_context=storage_context)

    # Step 4: Insert the nodes
    index.insert_nodes(nodes)

    # Step 5: Save to disk if needed
    if persist:
        index.storage_context.persist(persist_dir=PERSIST_DIR)

    return index, index.as_query_engine()



def load_or_build_index(documents: List[Document]):
    # Check if all required persisted files exist
    required_files = ["docstore.json", "vector_store.json", "index_store.json"]
    persist_files_exist = all(os.path.exists(os.path.join(PERSIST_DIR, f)) for f in required_files)

    if persist_files_exist:
        try:
            print("📦 Loading existing index from disk...")
            storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
            index = load_index_from_storage(storage_context)
            return index, index.as_query_engine()
        except Exception as e:
            print(f"⚠️ Failed to load index from storage: {e}")

    print("🛠️ Building new index...")
    return build_index_and_engine(documents)


def hash_file(filepath: str) -> str:
    hasher = hashlib.sha256()
    with open(filepath, 'rb') as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()

def get_files_snapshot(directory: str) -> Dict[str, str]:
    snapshot = {}
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            try:
                snapshot[filename] = hash_file(file_path)
            except Exception as e:
                print(f"❌ Failed to hash {filename}: {e}")
    return snapshot


def parse_query(user_query: str) -> Tuple[str | None, str]:
    # Match phrases like: "from sales.xlsx", "in document 'data.csv'", "regarding file data.xlsx"
    file_match = re.search(
        r"(?:from|in|of|regarding|about|open|for|in the document|in file)\s+[\"']?([\w\-\s]+\.(?:txt|pdf|docx|csv|xlsx?|xls))[\"']?",
        user_query, re.IGNORECASE
    )

    if file_match:
        file_name = file_match.group(1).strip()
        # Remove matched portion from the query
        stripped_query = user_query.replace(file_match.group(0), "").strip()
        return file_name, stripped_query or f"Summarize the contents of {file_name}"
    
    # Extra: if user query ends with a known file
    trailing_match = re.search(r"([\w\-\s]+\.(?:txt|pdf|docx|csv|xlsx?|xls))$", user_query.strip(), re.IGNORECASE)
    if trailing_match:
        file_name = trailing_match.group(1).strip()
        return file_name, f"Summarize the contents of {file_name}"

    return None, user_query


class FileChangeHandler(FileSystemEventHandler):
    def on_any_event(self, event):
        global documents, query_engine, file_snapshot
        if event.is_directory:
            return

        current_snapshot = get_files_snapshot(FILES_DIR)

        if current_snapshot != file_snapshot:
            with lock:
                print("\n🔁 File content changed. Reloading documents...")
                documents = load_documents(FILES_DIR)
                if documents:
                    _, query_engine = load_or_build_index(documents)
                    file_snapshot = current_snapshot
                    print("✅ Documents reloaded.")
                else:
                    print("⚠️ No supported documents found after reload.")
        else:
            print("ℹ️ File change detected but contents are the same. Skipping reload.")

def normalize(name: str) -> str:
    return name.strip().lower().replace(" ", "")

def main():
    global documents, query_engine, file_snapshot

    if not os.path.exists(FILES_DIR):
        os.makedirs(FILES_DIR)
        print(f"📂 Created '{FILES_DIR}' folder. Please add some documents and restart.")
        return

    hf_login(token="") # "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"

    system_prompt = (
        "You are a SuperLlama, A document summarization assistant which uses Meta LLama 3.2 LLM."
        "Your task is to answer questions based on the provided documents."
        "You can only use the information from the documents in the 'files/' directory."
        "If a question refers to a specific file, you should only answer based on that file."
        "If the question does not refer to a specific file, you should search all files."
        "You are not limited to answering questions only about the provided documents."
        "You can also answer general questions, but your primary focus is on the documents."
        "You can have to summarize the contents of a file if requested."
    )
    embed_model = get_embed_model()
    llm = get_llm(system_prompt)

    Settings.llm = llm
    Settings.embed_model = embed_model
    Settings.chunk_size = 1024
    Settings.context_window = 4096

    print("\n📚 Checking documents in 'files/' directory...")
    documents = load_documents(FILES_DIR)
    if not documents:
        print("⚠️ No supported documents found in 'files/' folder. Add .txt, .pdf, .docx, .csv, .xls, or .xlsx files.")
        return

    _, query_engine = load_or_build_index(documents)
    file_snapshot = get_files_snapshot(FILES_DIR)

    # Start file monitoring
    event_handler = FileChangeHandler()
    observer = Observer()
    observer.schedule(event_handler, path=FILES_DIR, recursive=False)
    observer.start()

    print("\n🤖 Q&A Assistant is ready! Type your question (or 'exit' to quit):")
    try:
        while True:
            user_query = input("\n\n").strip()
            if user_query.lower() in ['exit', 'quit']:
                print("👋 Goodbye!")
                break

            with lock:
                target_file, stripped_query = parse_query(user_query)

                if target_file:
                    matching_docs = [ doc for doc in documents if normalize(doc.metadata.get("file_name", "")) == normalize(target_file) ]
                    if matching_docs:
                        _, filtered_engine = build_index_and_engine(matching_docs)
                        response = filtered_engine.query(stripped_query)
                        if not str(response).strip():
                            response = f"🤖 I couldn't find an answer in '{target_file}'. Try rephrasing your question."
                    else:
                        response = f"⚠️ No document named '{target_file}' found in 'files/' directory."
                else:
                    response = query_engine.query(user_query)

            print(f"Assistant: {response}\n")

    finally:
        observer.stop()
        observer.join()

if __name__ == "__main__":
    main()
