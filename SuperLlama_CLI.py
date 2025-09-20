import os
import time
import torch
import re
import pandas as pd
from typing import Tuple, List, Dict
import fitz
from docx import Document as DocxDocument
from llama_index.core import Document, Settings, VectorStoreIndex
from llama_index.llms.huggingface import HuggingFaceLLM
from llama_index.embeddings.langchain import LangchainEmbedding
from langchain_huggingface import HuggingFaceEmbeddings
from huggingface_hub import login
from transformers import BitsAndBytesConfig
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from threading import Lock
import hashlib
import shutil
import logging

# Suppress verbose warnings from the transformers package
logging.getLogger("transformers").setLevel(logging.ERROR)

from llama_index.core import StorageContext, load_index_from_storage
from rich.console import Console
from rich.panel import Panel
from rich.rule import Rule
from rich.markdown import Markdown
from rich.prompt import Prompt

# --- Globals & Constants ---
PERSIST_DIR = "index_storage"
FILES_DIR = "files"
console = Console()
documents = []
query_engine = None
file_snapshot = {}
lock = Lock()


def dataframe_to_markdown(df: pd.DataFrame, filename: str) -> str:
    """Converts a pandas DataFrame to a formatted Markdown string for display."""
    MAX_ROWS = 50
    preview_df = df.head(MAX_ROWS)

    metadata = [
        f"### Table Preview from `{filename}`",
        f"- **Shape**: {df.shape[0]} rows × {df.shape[1]} columns",
        "- **Columns**:",
        *(f"  - `{col}`: {dtype}" for col, dtype in zip(df.columns, df.dtypes)),
        "",
        preview_df.to_markdown(index=False)
    ]
    return "\n".join(metadata)

# --- File Loaders ---
FILE_LOADERS = {
    '.txt': lambda p: open(p, 'r', encoding='utf-8').read(),
    '.pdf': lambda p: "".join(page.get_text() for page in fitz.open(p)),
    '.docx': lambda p: "\n".join(para.text for para in DocxDocument(p).paragraphs),
    '.csv': lambda p: dataframe_to_markdown(pd.read_csv(p), os.path.basename(p)),
    '.xls': lambda p: dataframe_to_markdown(pd.read_excel(p), os.path.basename(p)),
    '.xlsx': lambda p: dataframe_to_markdown(pd.read_excel(p), os.path.basename(p)),
}

# --- Core LlamaIndex & Model Functions ---
def hf_login(token: str):
    """Logs into Hugging Face Hub."""
    login(token=token)

def get_embed_model():
    """Initializes and returns the sentence-transformer embedding model."""
    return LangchainEmbedding(
        HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    )

def get_llm(system_prompt: str):
    """Initializes and returns the HuggingFace LLM with quantization."""
    quant_config = BitsAndBytesConfig(load_in_8bit=True)
    return HuggingFaceLLM(
        context_window=4096,
        max_new_tokens=256,
        system_prompt=system_prompt,
        tokenizer_name="meta-llama/Llama-3.2-1B-Instruct",
        model_name="meta-llama/Llama-3.2-1B-Instruct",
        device_map="auto",
        model_kwargs={
            "torch_dtype": torch.float16,
            "quantization_config": quant_config,
        }
    )

from llama_index.core.node_parser import SimpleNodeParser

def chunk_documents(documents: List[Document]) -> List[Document]:
    """Chunks a list of documents into smaller nodes."""
    parser = SimpleNodeParser.from_defaults(chunk_size=1024, chunk_overlap=200)
    nodes = []
    for doc in documents:
        try:
            nodes.extend(parser.get_nodes_from_documents([doc]))
        except Exception as e:
            console.print(f"[bold red]Failed to chunk document {doc.metadata.get('file_name', '')}: {e}[/bold red]")
    return nodes


def load_documents(directory: str) -> List[Document]:
    """Loads all supported documents from a given directory."""
    docs = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        ext = os.path.splitext(filename)[1].lower()
        if ext in FILE_LOADERS:
            try:
                content = FILE_LOADERS[ext](file_path)
                docs.append(Document(text=content, metadata={"file_name": filename}))
                console.print(f"[green]Loaded: {filename}[/green]")
            except Exception as e:
                console.print(f"[bold red]Failed to load {filename}: {e}[/bold red]")
        else:
            console.print(f"[yellow]Skipped unsupported file: {filename}[/yellow]")
    return docs


from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.core.schema import TextNode

def build_index_and_engine(documents: List[Document], persist: bool = True):
    """Builds a new VectorStoreIndex and query engine from documents."""
    nodes = chunk_documents(documents)
    storage_context = StorageContext.from_defaults()
    index = VectorStoreIndex.from_documents([], storage_context=storage_context)
    index.insert_nodes(nodes)
    if persist:
        index.storage_context.persist(persist_dir=PERSIST_DIR)
    return index, index.as_query_engine()


def load_or_build_index(documents: List[Document]):
    """Loads an index from disk if it exists, otherwise builds a new one."""
    required_files = ["docstore.json", "vector_store.json", "index_store.json"]
    if all(os.path.exists(os.path.join(PERSIST_DIR, f)) for f in required_files):
        try:
            console.print("[bold]Loading existing index from disk...[/bold]")
            storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
            index = load_index_from_storage(storage_context)
            return index, index.as_query_engine()
        except Exception as e:
            console.print(f"[bold yellow]Failed to load index from storage: {e}. Rebuilding...[/bold yellow]")

    console.print("[bold]Building new index...[/bold]")
    return build_index_and_engine(documents)

# --- File Monitoring & Query Parsing ---
def hash_file(filepath: str) -> str:
    """Generates a SHA256 hash for a given file."""
    hasher = hashlib.sha256()
    with open(filepath, 'rb') as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()

def get_files_snapshot(directory: str) -> Dict[str, str]:
    """Creates a dictionary snapshot of filenames and their hashes in a directory."""
    snapshot = {}
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            try:
                snapshot[filename] = hash_file(file_path)
            except Exception as e:
                console.print(f"[bold red]Failed to hash {filename}: {e}[/bold red]")
    return snapshot


def parse_query(user_query: str) -> Tuple[str | None, str]:
    """Parses a user query to extract a potential filename."""
    file_match = re.search(
        r"(?:from|in|of|regarding|about|open|for|in the document|in file)\s+[\"']?([\w\-\s]+\.(?:txt|pdf|docx|csv|xlsx?|xls))[\"']?",
        user_query, re.IGNORECASE
    )
    if file_match:
        file_name = file_match.group(1).strip()
        stripped_query = user_query.replace(file_match.group(0), "").strip()
        return file_name, stripped_query or f"Summarize the contents of {file_name}"
    
    trailing_match = re.search(r"([\w\-\s]+\.(?:txt|pdf|docx|csv|xlsx?|xls))$", user_query.strip(), re.IGNORECASE)
    if trailing_match:
        file_name = trailing_match.group(1).strip()
        return file_name, f"Summarize the contents of {file_name}"

    return None, user_query


class FileChangeHandler(FileSystemEventHandler):
    """Handles file system events to trigger document reloads."""
    def on_any_event(self, event):
        global documents, query_engine, file_snapshot
        if event.is_directory: return

        with lock:
            current_snapshot = get_files_snapshot(FILES_DIR)
            if current_snapshot != file_snapshot:
                console.print("\n[bold yellow]File change detected. Reloading documents...[/bold yellow]")
                with console.status("[bold green]Processing files...", spinner="dots"):
                    documents = load_documents(FILES_DIR)
                    if documents:
                        _, query_engine = load_or_build_index(documents)
                        file_snapshot = current_snapshot
                console.print("[bold green]Documents reloaded and index is up-to-date.[/bold green]")
            else:
                # This case is rare but handles metadata-only changes
                pass

def normalize(name: str) -> str:
    """Normalizes filenames for case-insensitive comparison."""
    return name.strip().lower().replace(" ", "")

def show_intro():
    """Clears the console and displays the welcome logo and introduction."""
    os.system('cls' if os.name == 'nt' else 'clear')
    art = r"""
  █████████                                          █████       ████                                     
 ███▒▒▒▒▒███                                        ▒▒███       ▒▒███                                     
▒███    ▒▒▒  █████ ████ ████████   ██████  ████████  ▒███        ▒███   ██████   █████████████    ██████  
▒▒█████████ ▒▒███ ▒███ ▒▒███▒▒███ ███▒▒███▒▒███▒▒███ ▒███        ▒███  ▒▒▒▒▒███ ▒▒███▒▒███▒▒███  ▒▒▒▒▒███ 
 ▒▒▒▒▒▒▒▒███ ▒███ ▒███  ▒███ ▒███▒███████  ▒███ ▒▒▒  ▒███        ▒███   ███████  ▒███ ▒███ ▒███   ███████ 
 ███    ▒███ ▒███ ▒███  ▒███ ▒███▒███▒▒▒   ▒███      ▒███      █ ▒███  ███▒▒███  ▒███ ▒███ ▒███  ███▒▒███ 
▒▒█████████  ▒▒████████ ▒███████ ▒▒██████  █████     ███████████ █████▒▒████████ █████▒███ █████▒▒████████
 ▒▒▒▒▒▒▒▒▒    ▒▒▒▒▒▒▒▒  ▒███▒▒▒   ▒▒▒▒▒▒  ▒▒▒▒▒     ▒▒▒▒▒▒▒▒▒▒▒ ▒▒▒▒▒  ▒▒▒▒▒▒▒▒ ▒▒▒▒▒ ▒▒▒ ▒▒▒▒▒  ▒▒▒▒▒▒▒▒ 
                        ▒███                                                                              
                        █████                                                                             
                       ▒▒▒▒▒                                                                              
    """
    console.print(art, style="bold blue")
    
    welcome_message = (
        "Welcome to [bold][blue]SuperLlama[/blue][/bold], your personal helper to boost productivity, powered by [bold][cyan]Meta Llama 3.2 1B[/cyan][/bold].\n"
        "Developed by [bold green]Sumit Saha[/bold green] | [bold green]GitHub: SumitBana[/bold green]\n\n"
        "[italic red]This Project is a part of my Undergrad Engineering Curriculum.[/italic red]"
        "\n\n[bold]Instructions:[/bold]\n"
        "1. Place your documents[red] BEFORE [/red]in the 'files/' directory.\n"
        "2. Ask questions about your documents using natural language.\n"
        "3. Use [bold]\"info\"[/bold] to see project details.\n"
        "4. Use [bold]\"cls\"[/bold] to clear the screen and see this message again.\n"
        "5. Use [bold]\"exit\"[/bold] or [bold]\"quit\"[/bold] to end the session."
    )
    console.print(Panel(welcome_message, title="[cyan]Introduction[/cyan]", border_style="cyan", title_align="center"))

def show_info():
    """Displays information about the project."""
    info_message = (

        "[blue]SuperLlama[/blue] is a local, document-aware chatbot designed to help you quickly find information within your files. "
        "It's powered by the compact yet powerful [bold][cyan]Meta Llama 3.2 1B[/cyan][/bold] model, running entirely on your machine.\n\n"
        "Developed by [bold green]Sumit Saha[/bold green] ([bold]GitHub: SumitBana[/bold]).\n\n"

        "Here's how it works under the hood:\n\n"
        "[bold][yellow]Document Chunking[/yellow][/bold]\n"
        "To handle large documents, the system first breaks them down into smaller, digestible pieces called chunks. "
        "This is crucial because language models have a limited 'context window' (the amount of text they can process at once).\n"
        "• [bold]Chunk Size:[/bold] Each chunk is set to 1024 characters.\n"
        "• [bold]Overlap:[/bold] To ensure that context isn't lost at the boundaries, consecutive chunks share 200 characters of text.\n\n"

        "[bold][yellow]Embedding Generation[/yellow][/bold]\n"
        "Next, these text chunks are converted into a machine-understandable format: numerical vectors. "
        "This process, called embedding, captures the semantic meaning of the text.\n"
        "• [bold]Model:[/bold] The script uses the [blue]sentence-transformers/all-mpnet-base-v2[/blue] model to generate these embeddings.\n"
        "• [bold]Outcome:[/bold] Each chunk becomes a vector, essentially a point on a high-dimensional 'meaning map.' "
        "Chunks with similar topics will have vectors that are close to each other.\n\n"

        "[bold][yellow]Vector Indexing & Retrieval[/yellow][/bold]\n"
        "All the generated vectors are stored and organized in a searchable database called a VectorStoreIndex."
        "This index is the core of the retrieval system.\n"
        "• [bold]Querying:[/bold] When you ask a question, your query is also converted into a vector using the same embedding model.\n"
        "• [bold]Semantic Search:[/bold] The system then searches the VectorStoreIndex to find the text chunks whose vectors are most similar to your query's vector. "
        "This is how it finds the most relevant information to answer your question.\n\n"

        "[bold][yellow]LLM-Powered Answering[/yellow][/bold]\n"
        "Finally, the most relevant chunks retrieved from the index are passed as context to the Large Language Model (LLM).\n"
        "• [bold]The Prompt:[/bold] The LLM receives your original question along with the retrieved text chunks.\n"
        "• [bold]Model:[/bold] The [blue]meta-llama/Llama-3.2-1B-Instruct[/blue] model synthesizes this information to generate a final, human-readable answer that is grounded in the provided documents."
    )
    console.print(Panel(info_message, title="[cyan]Project Information[/cyan]", border_style="cyan", title_align="center"))

# --- Main Application Logic ---
def main():
    global documents, query_engine, file_snapshot

    # --- Check for and remove existing index directory at the start ---
    if os.path.exists(PERSIST_DIR):
        console.print(f"[bold yellow]Found existing index directory '{PERSIST_DIR}'. Removing it...[/bold yellow]")
        try:
            shutil.rmtree(PERSIST_DIR)
            console.print(f"[green]Successfully removed '{PERSIST_DIR}'.[/green]")
        except Exception as e:
            console.print(f"[bold red]Error removing directory '{PERSIST_DIR}': {e}[/bold red]")
            console.print("[bold red]Please remove it manually and restart the application.[/bold red]")
            return

    if not os.path.exists(FILES_DIR):
        os.makedirs(FILES_DIR)
        console.print(f"[bold yellow]Created '{FILES_DIR}' folder. Please add some documents and restart.[/bold yellow]")
        return

    # --- CLI Interface ---
    show_intro()

    # --- Initialization ---
    with console.status("[bold green]Initializing models and settings...", spinner="dots"):
        hf_login(token="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")

        system_prompt = (
            "You are SuperLlama, a document summarization assistant powered by Meta Llama 3.2. "
            "Your primary task is to answer questions based on the documents provided in the 'files/' directory. "
            "If a question refers to a specific file, answer only based on that file. "
            "If not, search all available files. You can also answer general questions, "
            "but always prioritize document-based information. Summarize file contents when requested."
        )
        Settings.llm = get_llm(system_prompt)
        Settings.embed_model = get_embed_model()
        Settings.chunk_size = 1024
        Settings.context_window = 4096

    with console.status("[bold green]Scanning and loading documents...", spinner="dots"):
        documents = load_documents(FILES_DIR)
    
    if not documents:
        console.print("[bold red]No supported documents found. Please add files to the 'files/' folder and restart.[/bold red]")
        return
    
    with console.status("[bold green]Building or loading knowledge index...", spinner="dots"):
        _, query_engine = load_or_build_index(documents)
        file_snapshot = get_files_snapshot(FILES_DIR)

    console.print("\n[bold green]Setup complete! You can now ask questions about your documents.[/bold green]")

    # --- File Monitoring & Main Loop ---
    event_handler = FileChangeHandler()
    observer = Observer()
    observer.schedule(event_handler, path=FILES_DIR, recursive=False)
    observer.start()

    try:
        while True:
            console.print(Rule(style="dim white"))
            user_query = Prompt.ask("[bold green]>").strip()
            
            if user_query.lower() in ['exit', 'quit']:
                console.print("[bold yellow]Goodbye![/bold yellow]")
                break
            
            if user_query.lower() == 'cls':
                show_intro()
                continue
            
            if user_query.lower() == 'info':
                show_info()
                continue

            response_str = ""
            with lock:
                with console.status("[bold cyan]Thinking...", spinner="earth"):
                    target_file, stripped_query = parse_query(user_query)

                    if target_file:
                        matching_docs = [ doc for doc in documents if normalize(doc.metadata.get("file_name", "")) == normalize(target_file) ]
                        if matching_docs:
                            _, filtered_engine = build_index_and_engine(matching_docs, persist=False)
                            response = filtered_engine.query(stripped_query)
                            response_str = str(response).strip() or f"I couldn't find a clear answer in '{target_file}'. Try rephrasing."
                        else:
                            response_str = f"No document named '{target_file}' found in the 'files/' directory."
                    else:
                        response = query_engine.query(user_query)
                        response_str = str(response)

            console.print(Panel(Markdown(response_str), border_style="blue", title="[bold blue]SuperLlama[/bold blue]", title_align="left"))

    finally:
        observer.stop()
        observer.join()
        console.print("\n" + "-"*20)
        console.print("[bold yellow]Cleaning up index files...[/bold yellow]")
        if os.path.exists(PERSIST_DIR):
            try:
                shutil.rmtree(PERSIST_DIR)
                console.print(f"[green]Successfully deleted '{PERSIST_DIR}' directory.[/green]")
            except Exception as e:
                console.print(f"[bold red]Error cleaning up '{PERSIST_DIR}': {e}[/bold red]")
        console.print("[bold yellow]Exiting application.[/bold yellow]")


if __name__ == "__main__":
    main()
